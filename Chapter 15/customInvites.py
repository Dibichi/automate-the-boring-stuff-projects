#! python3
#customInvites - pulls text from a text file by line and creates a unique word document with that text file
import docx
import os

userDir = input('What directory is the text file in\n')
os.chdir(userDir)

userFile = input('What is the text file called\n')

doc = docx.Document('invitation.docx')
txtFile = open(userFile)

content = txtFile.readlines()

for line in content:
    name = line.strip()

    doc.add_paragraph('It would be a pleasure to have the company of', style = 'Brush Script Std')
    doc.add_paragraph(name, style = 'Segoe UI')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of', style = 'Brush Script Std')
    doc.add_paragraph('April 1st', style = 'Segoe UI')
    doc.add_paragraph("at 7 o'clock", style = 'Brush Script Std')

    doc.add_page_break()

doc.save('invitation.docx')

print('Done - file name is invitation.docx')

